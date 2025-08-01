import os
import json
import logging
import mimetypes
from typing import Dict, Any, Optional, List
import asyncio
from pathlib import Path
import base64

# File processing libraries
import pandas as pd
from PIL import Image
import PyPDF2
import docx
import openpyxl
import csv
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles processing of various file types"""
    
    def __init__(self):
        self.supported_types = {
            # Text files
            'text/plain': self._process_text,
            'text/csv': self._process_csv,
            'application/json': self._process_json,
            'application/xml': self._process_xml,
            'text/xml': self._process_xml,
            
            # Documents
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
            
            # Spreadsheets
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'application/vnd.ms-excel': self._process_xls,
            
            # Images
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
            'image/gif': self._process_image,
            'image/bmp': self._process_image,
            'image/webp': self._process_image,
            
            # Code files
            'text/javascript': self._process_code,
            'text/x-python': self._process_code,
            'application/x-python-code': self._process_code,
            'text/html': self._process_code,
            'text/css': self._process_code,
        }
    
    async def process_file(self, file_path: str, content_type: Optional[str] = None) -> Dict[str, Any]:
        """Process file and extract content/metadata"""
        try:
            # Determine content type if not provided
            if not content_type:
                content_type, _ = mimetypes.guess_type(file_path)
            
            # Get file info
            file_info = self._get_file_info(file_path)
            
            # Process based on content type
            if content_type in self.supported_types:
                processor = self.supported_types[content_type]
                content = await processor(file_path)
            else:
                # Try to process as text if unknown type
                try:
                    content = await self._process_text(file_path)
                except:
                    content = {"error": f"Unsupported file type: {content_type}"}
            
            return {
                "file_info": file_info,
                "content_type": content_type,
                "processed_content": content,
                "processing_time": file_info.get("processing_time"),
                "success": "error" not in content
            }
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return {
                "error": str(e),
                "file_path": file_path,
                "success": False
            }
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            file_stat = os.stat(file_path)
            path_obj = Path(file_path)
            
            return {
                "filename": path_obj.name,
                "extension": path_obj.suffix,
                "size_bytes": file_stat.st_size,
                "size_mb": round(file_stat.st_size / (1024 * 1024), 2),
                "created_time": file_stat.st_ctime,
                "modified_time": file_stat.st_mtime,
                "is_binary": self._is_binary_file(file_path)
            }
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {"error": str(e)}
    
    def _is_binary_file(self, file_path: str) -> bool:
        """Check if file is binary"""
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                return b'\0' in chunk
        except:
            return True
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Basic text analysis
            lines = content.split('\n')
            words = content.split()
            
            return {
                "content": content,
                "line_count": len(lines),
                "word_count": len(words),
                "character_count": len(content),
                "preview": content[:500] + "..." if len(content) > 500 else content
            }
            
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return {"error": str(e)}
    
    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Basic analysis
            analysis = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "data_types": df.dtypes.to_dict(),
                "null_counts": df.isnull().sum().to_dict(),
                "preview": df.head(10).to_dict('records'),
                "summary_stats": df.describe().to_dict() if df.select_dtypes(include='number').empty == False else {}
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {e}")
            return {"error": str(e)}
    
    async def _process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            return {
                "content": data,
                "type": type(data).__name__,
                "size": len(str(data)),
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "length": len(data) if isinstance(data, (list, dict)) else None
            }
            
        except Exception as e:
            logger.error(f"Error processing JSON file: {e}")
            return {"error": str(e)}
    
    async def _process_xml(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def xml_to_dict(element):
                result = {}
                if element.text and element.text.strip():
                    result['text'] = element.text.strip()
                
                for child in element:
                    child_data = xml_to_dict(child)
                    if child.tag in result:
                        if not isinstance(result[child.tag], list):
                            result[child.tag] = [result[child.tag]]
                        result[child.tag].append(child_data)
                    else:
                        result[child.tag] = child_data
                
                if element.attrib:
                    result['@attributes'] = element.attrib
                
                return result
            
            content = xml_to_dict(root)
            
            return {
                "root_tag": root.tag,
                "content": content,
                "namespace": root.tag.split('}')[0][1:] if '}' in root.tag else None
            }
            
        except Exception as e:
            logger.error(f"Error processing XML file: {e}")
            return {"error": str(e)}
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                text_content = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
                
                metadata = pdf_reader.metadata
                
                return {
                    "page_count": len(pdf_reader.pages),
                    "text_content": text_content,
                    "metadata": {
                        "title": metadata.get('/Title', '') if metadata else '',
                        "author": metadata.get('/Author', '') if metadata else '',
                        "subject": metadata.get('/Subject', '') if metadata else '',
                        "creator": metadata.get('/Creator', '') if metadata else ''
                    },
                    "word_count": len(text_content.split()),
                    "preview": text_content[:1000] + "..." if len(text_content) > 1000 else text_content
                }
                
        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            return {"error": str(e)}
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            text_content = '\n'.join(full_text)
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return {
                "text_content": text_content,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "tables": tables_data,
                "word_count": len(text_content.split()),
                "preview": text_content[:1000] + "..." if len(text_content) > 1000 else text_content
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            return {"error": str(e)}
    
    async def _process_doc(self, file_path: str) -> Dict[str, Any]:
        """Process DOC files (legacy Word format)"""
        try:
            # For DOC files, we'd need python-docx2txt or similar
            # For now, return basic info
            return {
                "error": "DOC format not fully supported. Please convert to DOCX format.",
                "suggestion": "Use a more modern format like DOCX for better processing"
            }
            
        except Exception as e:
            logger.error(f"Error processing DOC file: {e}")
            return {"error": str(e)}
    
    async def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Process Excel XLSX files"""
        try:
            # Read with pandas
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist(),
                    "preview": df.head(5).to_dict('records'),
                    "data_types": df.dtypes.to_dict()
                }
            
            return {
                "sheet_names": excel_file.sheet_names,
                "sheet_count": len(excel_file.sheet_names),
                "sheets_data": sheets_data
            }
            
        except Exception as e:
            logger.error(f"Error processing XLSX file: {e}")
            return {"error": str(e)}
    
    async def _process_xls(self, file_path: str) -> Dict[str, Any]:
        """Process Excel XLS files"""
        try:
            # Similar to XLSX but for older format
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist(),
                    "preview": df.head(5).to_dict('records')
                }
            
            return {
                "sheet_names": excel_file.sheet_names,
                "sheet_count": len(excel_file.sheet_names),
                "sheets_data": sheets_data
            }
            
        except Exception as e:
            logger.error(f"Error processing XLS file: {e}")
            return {"error": str(e)}
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files"""
        try:
            with Image.open(file_path) as img:
                # Get image info
                width, height = img.size
                format_type = img.format
                mode = img.mode
                
                # Convert to base64 for preview (if small enough)
                file_size = os.path.getsize(file_path)
                base64_data = None
                
                if file_size < 1024 * 1024:  # Less than 1MB
                    with open(file_path, 'rb') as f:
                        base64_data = base64.b64encode(f.read()).decode('utf-8')
                
                return {
                    "width": width,
                    "height": height,
                    "format": format_type,
                    "mode": mode,
                    "has_transparency": img.mode in ('RGBA', 'LA'),
                    "base64_preview": base64_data,
                    "megapixels": round((width * height) / 1000000, 2)
                }
                
        except Exception as e:
            logger.error(f"Error processing image file: {e}")
            return {"error": str(e)}
    
    async def _process_code(self, file_path: str) -> Dict[str, Any]:
        """Process code files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            # Basic code analysis
            non_empty_lines = [line for line in lines if line.strip()]
            comment_lines = []
            
            # Detect comments based on file extension
            extension = Path(file_path).suffix.lower()
            if extension in ['.py']:
                comment_lines = [line for line in lines if line.strip().startswith('#')]
            elif extension in ['.js', '.css', '.java', '.cpp', '.c']:
                comment_lines = [line for line in lines if line.strip().startswith('//') or '/*' in line]
            elif extension in ['.html', '.xml']:
                comment_lines = [line for line in lines if '<!--' in line or '-->' in line]
            
            return {
                "content": content,
                "language": self._detect_language(extension),
                "total_lines": len(lines),
                "code_lines": len(non_empty_lines),
                "comment_lines": len(comment_lines),
                "blank_lines": len(lines) - len(non_empty_lines),
                "preview": content[:1000] + "..." if len(content) > 1000 else content,
                "functions": self._extract_functions(content, extension),
                "imports": self._extract_imports(content, extension)
            }
            
        except Exception as e:
            logger.error(f"Error processing code file: {e}")
            return {"error": str(e)}
    
    def _detect_language(self, extension: str) -> str:
        """Detect programming language from file extension"""
        language_map = {
            '.py': 'Python',
            '.js': 'JavaScript',
            '.html': 'HTML',
            '.css': 'CSS',
            '.java': 'Java',
            '.cpp': 'C++',
            '.c': 'C',
            '.php': 'PHP',
            '.rb': 'Ruby',
            '.go': 'Go',
            '.rs': 'Rust',
            '.ts': 'TypeScript',
            '.jsx': 'React JSX',
            '.tsx': 'React TSX',
            '.vue': 'Vue',
            '.sql': 'SQL'
        }
        return language_map.get(extension.lower(), 'Unknown')
    
    def _extract_functions(self, content: str, extension: str) -> List[str]:
        """Extract function names from code"""
        functions = []
        lines = content.split('\n')
        
        try:
            if extension == '.py':
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('def ') and '(' in stripped:
                        func_name = stripped[4:stripped.index('(')].strip()
                        functions.append(func_name)
                    elif stripped.startswith('class ') and ':' in stripped:
                        class_name = stripped[6:stripped.index(':')].strip()
                        functions.append(f"class {class_name}")
            
            elif extension in ['.js', '.ts']:
                for line in lines:
                    stripped = line.strip()
                    if 'function ' in stripped and '(' in stripped:
                        start = stripped.find('function ') + 9
                        end = stripped.find('(', start)
                        if end > start:
                            func_name = stripped[start:end].strip()
                            functions.append(func_name)
                    elif '=>' in stripped and '=' in stripped:
                        # Arrow functions
                        parts = stripped.split('=')
                        if len(parts) > 1:
                            func_name = parts[0].strip().replace('const', '').replace('let', '').replace('var', '').strip()
                            functions.append(func_name)
            
        except Exception:
            pass
        
        return functions
    
    def _extract_imports(self, content: str, extension: str) -> List[str]:
        """Extract import statements from code"""
        imports = []
        lines = content.split('\n')
        
        try:
            if extension == '.py':
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('import ') or stripped.startswith('from '):
                        imports.append(stripped)
            
            elif extension in ['.js', '.ts']:
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('import ') or stripped.startswith('const ') and 'require(' in stripped:
                        imports.append(stripped)
            
        except Exception:
            pass
        
        return imports
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats"""
        return {
            "text": [".txt", ".md", ".csv", ".json", ".xml"],
            "documents": [".pdf", ".docx", ".doc"],
            "spreadsheets": [".xlsx", ".xls"],
            "images": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"],
            "code": [".py", ".js", ".html", ".css", ".java", ".cpp", ".c", ".php", ".rb", ".go", ".rs", ".ts", ".jsx", ".tsx", ".vue", ".sql"]
        }
    
    async def batch_process_files(self, file_paths: List[str]) -> Dict[str, Any]:
        """Process multiple files in batch"""
        results = {}
        errors = []
        
        for file_path in file_paths:
            try:
                result = await self.process_file(file_path)
                results[file_path] = result
            except Exception as e:
                error_info = {
                    "file_path": file_path,
                    "error": str(e)
                }
                errors.append(error_info)
                logger.error(f"Batch processing error for {file_path}: {e}")
        
        return {
            "processed_files": len(results),
            "errors": len(errors),
            "results": results,
            "error_details": errors
        }
    
    async def analyze_directory(self, directory_path: str) -> Dict[str, Any]:
        """Analyze all files in a directory"""
        try:
            dir_path = Path(directory_path)
            if not dir_path.exists() or not dir_path.is_dir():
                return {"error": "Directory does not exist or is not a directory"}
            
            file_analysis = {
                "total_files": 0,
                "total_size_mb": 0,
                "file_types": {},
                "largest_files": [],
                "supported_files": [],
                "unsupported_files": []
            }
            
            all_files = []
            for file_path in dir_path.rglob('*'):
                if file_path.is_file():
                    all_files.append(file_path)
            
            file_analysis["total_files"] = len(all_files)
            
            for file_path in all_files:
                try:
                    file_stat = file_path.stat()
                    file_size_mb = file_stat.st_size / (1024 * 1024)
                    file_analysis["total_size_mb"] += file_size_mb
                    
                    # Track file types
                    extension = file_path.suffix.lower()
                    if extension in file_analysis["file_types"]:
                        file_analysis["file_types"][extension] += 1
                    else:
                        file_analysis["file_types"][extension] = 1
                    
                    # Track largest files
                    file_info = {
                        "path": str(file_path),
                        "size_mb": round(file_size_mb, 2),
                        "extension": extension
                    }
                    file_analysis["largest_files"].append(file_info)
                    
                    # Check if supported
                    content_type, _ = mimetypes.guess_type(str(file_path))
                    if content_type in self.supported_types or extension in ['.py', '.js', '.html', '.css']:
                        file_analysis["supported_files"].append(str(file_path))
                    else:
                        file_analysis["unsupported_files"].append(str(file_path))
                
                except Exception as e:
                    logger.warning(f"Error analyzing file {file_path}: {e}")
            
            # Sort largest files
            file_analysis["largest_files"].sort(key=lambda x: x["size_mb"], reverse=True)
            file_analysis["largest_files"] = file_analysis["largest_files"][:10]  # Top 10
            
            file_analysis["total_size_mb"] = round(file_analysis["total_size_mb"], 2)
            
            return file_analysis
            
        except Exception as e:
            logger.error(f"Error analyzing directory: {e}")
            return {"error": str(e)}
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up temporary uploaded files"""
        try:
            uploads_dir = Path("uploads")
            if not uploads_dir.exists():
                return
            
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            cleaned_count = 0
            for file_path in uploads_dir.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        cleaned_count += 1
                        logger.info(f"Cleaned up old file: {file_path}")
            
            logger.info(f"Cleaned up {cleaned_count} old files")
            
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")
    
    async def extract_text_for_ai(self, file_path: str) -> str:
        """Extract text content specifically for AI processing"""
        try:
            content_type, _ = mimetypes.guess_type(file_path)
            
            if content_type == 'application/pdf':
                result = await self._process_pdf(file_path)
                return result.get('text_content', '')
            
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                result = await self._process_docx(file_path)
                return result.get('text_content', '')
            
            elif content_type == 'text/csv':
                result = await self._process_csv(file_path)
                # Convert CSV data to readable text
                if 'preview' in result:
                    preview_text = "CSV Data Preview:\n"
                    for row in result['preview']:
                        preview_text += str(row) + "\n"
                    return preview_text
                return str(result)
            
            elif content_type == 'application/json':
                result = await self._process_json(file_path)
                return json.dumps(result.get('content', {}), indent=2)
            
            elif content_type and content_type.startswith('text/'):
                result = await self._process_text(file_path)
                return result.get('content', '')
            
            else:
                # Try as text file
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Error extracting text for AI: {e}")
            return f"Error extracting text from file: {str(e)}"